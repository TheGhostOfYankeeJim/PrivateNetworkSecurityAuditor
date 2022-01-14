# Obj-Ori Programming attempt

import sys
import subprocess
import os
import re
import time

import docx 
from datetime import date
from shutil import copyfile

import netifaces as networkInterface
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *

class MatchLineEdit(QLineEdit):
	def focusInEvent(self, e):
		try:
			subprocess.Popen(["matchbox-keyboard"])
		except FileNotFoundError:
			pass

	def focusOutEvent(self,e):
		subprocess.Popen(["kill","matchbox-keyboard"])
 
 
class App(QMainWindow):
	def __init__(self):
		super().__init__()
		self.title = 'Personal Network Security Assistant'
		self.left = 0
		self.top = 0
		self.width = 800
		self.height = 400
		self.statusBarMessage = "Welcome to the PNSA"
		self.currentSprite = 'logoTest.png'
		self.saveLocationDir = ""
		self.ipAddressUsed = ""
		self.versionNumber = 1.0
		self.probablyRouter = ""
		
		print(QStyleFactory.keys())
		
		self.initUI()
		
	def initUI(self):
		self.setWindowTitle(self.title)
		self.setGeometry(self.left, self.top, self.width, self.height)
		self.statusBar().showMessage(self.statusBarMessage)
		
		
		print(self.saveLocationDir)
		
		# Set up the buttons
		btn1 = QPushButton('Start', self)
		btn1.setToolTip('Click me to start the test')
		btn1.resize(100,100)
		btn1.move(0,0)
		btn1.clicked.connect(self.btn1_clicked)
		
		btn2 = QPushButton('Config Wi-Fi', self)
		btn2.setToolTip('Click me to pause tests')
		btn2.resize(100,100)
		btn2.move(100,0)
		btn2.clicked.connect(self.btn2_clicked)
		
		btn3 = QPushButton('Report Out', self)
		btn3.setToolTip('Click me to stop tests')
		btn3.resize(100,100)
		btn3.move(200,0)
		btn3.clicked.connect(self.btn3_clicked)
		
		
		btn4 = QPushButton('Update', self)
		btn4.setToolTip('Click me to update PNSA')
		btn4.resize(100,100)
		btn4.move(300,0)
		btn4.clicked.connect(self.btn4_clicked)
		
		
		btn5 = QPushButton('About', self)
		btn5.setToolTip('Quick Overview of PNSA')
		btn5.resize(100,100)
		btn5.move(400,0)
		btn5.clicked.connect(self.btn5_clicked)
		
		
		btn6 = QPushButton('Quit/PowerOff', self)
		btn6.setToolTip('Kills the Application')
		btn6.resize(100,100)
		btn6.move(500,0)
		btn6.clicked.connect(self.btn6_clicked)
		
		
		spriteHelper = QLabel(self)
		spriteHelper.move(600,0)
		spriteHelper.setPixmap(QPixmap(self.currentSprite))
		spriteHelper.resize(200,300)
		
		
		#Page One Layout
		self.pageOneTitle = QLabel(self)
		self.pageOneTitle.move(0,100)
		self.pageOneTitle.setText("<b>Initial Setup Screen</b>")
		self.pageOneTitle.resize(120,20)
		
		self.userNameLabel = QLabel(self)
		self.userNameLabel.move(0,125)
		self.userNameLabel.setText("What is your name?")
		self.userNameLabel.resize(120,20)
		
		self.nameInput = MatchLineEdit(self)
		self.nameInput.move(0,145)
		self.nameInput.resize(200,32)
		self.nameInputText = ""
		
		self.wifiSSID = QLabel(self)
		self.wifiSSID.move(0,175)
		self.wifiSSID.setText("Enter Wi-Fi Name")
		self.wifiSSID.resize(120,20)
		
		self.wifiName = MatchLineEdit(self)
		self.wifiName.move(0,195)
		self.wifiName.resize(200,32)
		self.wifiName.setEchoMode(0)
		self.wifiNameText = ""
		
		
		self.wifiKey = QLabel(self)
		self.wifiKey.move(0,224)
		self.wifiKey.setText("Enter Wi-Fi Password")
		self.wifiKey.resize(120,20)
		
		self.wifiPassword = MatchLineEdit(self)
		self.wifiPassword.move(0,244)
		self.wifiPassword.resize(200,32)
		self.wifiPassword.setEchoMode(3)
		self.wifiPasswordText = ""
		
		#self.connectionLabel = QLabel(self)
		#self.connectionLabel.move(200,125)
		#self.connectionLabel.setText("Select Connection Type")
		#self.connectionLabel.resize(130,20)

		
		#self.comboConnection = QComboBox(self)
		#self.comboConnection.addItem("Wi-Fi")
		#self.comboConnection.addItem("Ethernet")
		#self.comboConnection.move(200,145)
		#self.comboConnection.resize(130,30)
		
		self.reportLabel = QLabel(self)
		self.reportLabel.move(200,125)
		self.reportLabel.setText("Select Report Type")
		self.reportLabel.resize(130,20)
		
		self.comboReport = QComboBox(self)
		self.comboReport.addItem("PDF")
		self.comboReport.addItem("Docx")
		self.comboReport.addItem("RTF")
		self.comboReport.addItem("TXT")
		self.comboReport.move(200,145)
		self.comboReport.resize(130,30)
	
		self.attackLabel = QLabel(self)
		self.attackLabel.move(350,125)
		self.attackLabel.setText("Test Selection")
		self.attackLabel.resize(130,20)
		
		self.cbWifi = QCheckBox('Wi-Fi Analysis', self)
		self.cbWifi.move(350, 145)
		self.cbWifi.resize(self.cbWifi.sizeHint())
		
		self.passAnal = QCheckBox('Password Analysis', self)
		self.passAnal.move(350, 165)
		self.passAnal.resize(self.passAnal.sizeHint())
		
		self.assDiscover = QCheckBox('Asset Discovery', self)
		self.assDiscover.move(350, 185)
		self.assDiscover.resize(self.assDiscover.sizeHint())

		self.portScan = QCheckBox('Port Scan', self)
		self.portScan.move(350, 205)
		self.portScan.resize(self.portScan.sizeHint())

		self.dhcpAttack = QCheckBox('DHCP Starvation', self)
		self.dhcpAttack.move(350, 225)
		self.dhcpAttack.resize(self.dhcpAttack.sizeHint())

		#self.arpAttack = QCheckBox('ARP Poisoning', self)
		#self.arpAttack.move(350, 245)
		#self.arpAttack.resize(self.arpAttack.sizeHint())
		
		
		self.saveLocation = QPushButton('Choose Save', self)
		self.saveLocation.setToolTip('Click to choose save location')
		self.saveLocation.resize(100,100)
		self.saveLocation.move(0,280)
		self.saveLocation.clicked.connect(self.saveLocation_clicked)
				
	@pyqtSlot()
	def btn1_clicked(self):
		
		self.statusBarMessage = "Are you sure?"
		self.statusBar().showMessage(self.statusBarMessage)
		
		reply = QMessageBox.question(self, 'Message', "Are you sure you want to start?", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)		
		
		if reply == QMessageBox.Yes:
			self.nameInputText = self.nameInput.text()
			self.wifiNameText = self.wifiName.text()
			self.wifiPasswordText = self.wifiPassword.text()
			
			self.testingSequence()
		else:
			pass
			
		# 0 means unchecked, 2 means checked
		#print(self.portScan.checkState())
		
		
		
	def btn2_clicked(self):
		
		self.statusBarMessage = "Configuring Wi-Fi"
		self.statusBar().showMessage(self.statusBarMessage)
		
		command = "nm-connection-editor"
		
		os.system(command)
		
	def btn3_clicked(self):
		
		QMessageBox.about(self, "Report Written", "The report has been exported to your save location")
		self.writeReport()
		
	def btn4_clicked(self):
		
		QMessageBox.about(self, "Update", "The PNSA is up to date")
	
	def btn5_clicked(self):
		message = """ Version 1.0 \n Author: Ian Briley \n This is a Proof Of Concept Application \n Please do not use this maliciously """
		QMessageBox.about(self, "About", message)
		
		
	
	def saveLocation_clicked(self):
		self.saveLocationDir = str(QFileDialog.getExistingDirectory(self, "Select Directory"))
		print(self.saveLocationDir)

	def btn6_clicked(self):
		reply = QMessageBox.question(self, 'Message', "Are you sure to quit?", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)		
		
		if reply == QMessageBox.Yes:
			sys.exit()
		else:
			print("Why did you click this then?")
	
	def writeReport(self):
		selectedType = self.comboReport.currentText()
		print("Write Report in: " + selectedType) 
		
		todaysDate = str(date.today())

		pnsaDoc = docx.Document()

		pnsaDoc.add_paragraph("PNSA Report", 'Title')
		
		pnsaDoc.add_paragraph(todaysDate)
		
		pnsaDoc.add_paragraph("Created by: Ian Briley")
		
		pnsaDoc.add_paragraph("CAPSTONE PROJECT DEMO REPORT WRITER")
		
		
		# Correct way to do this, book was wrong
		pnsaDoc.add_page_break()
		
		aboutPNSATitle = pnsaDoc.add_heading("About The PNSA:", 1)
		
		#I know this is an ugly line, will figure out how to do a multilines
		aboutTextBody = """	The Personal Network Security Assistant (PNSA) is a prototype proof of concept portable pen-testing device with Ethernet/WiFi connection for home and small business use. The PNSA consists a touchscreen display, and a custom made easy to use GUI."""
		aboutTextBody2 = """	The PNSA is responsible for conducting a standard suite of penetration tests and vulnerability assessments that would be best suited for a home or small business network. The tests currently consist of port-scan of the network, password analysis, Wi-Fi Analysis, asset discovery, DHCP Starvation, and Password Security Tool. The purpose of this PNSA prototype is to help home users and small businesses test, understand, and strengthen their networks at a basic level."""
		
		aboutPNSA = pnsaDoc.add_paragraph(aboutTextBody)
		aboutPNSA = pnsaDoc.add_paragraph(aboutTextBody2)
		
		purpose = pnsaDoc.add_heading("Purpose of Reporting:", 1)
		
		purposeOfPNSA = """	The primary purpose of conducting tests and analyzing the results to is be made aware of your network and any changes that are made over time. Without knowing your network, you can not defend it properly. Knowledge in this case is power."""
		purposeOfPNSA2 = """	Additionally, this reporting portion of the PNSA serves as a great start for users who want to learn more about their network without having to have an extensive network knowledge. Finally, at the very least, this provides a great report."""
		purposeOfPNSA3 = """	Finally, at the very least, this provides a great report for non tech savvy users to provide to a computer consulatation agent, and or their trusty friend that 'knows a lot about computers'. What you do with the knowledge gain is ultimately up to you, the PNSA is simply provides the means to get you that information."""
		
		purposePNSA = pnsaDoc.add_paragraph(purposeOfPNSA)
		purposePNSA = pnsaDoc.add_paragraph(purposeOfPNSA2)
		purposePNSA = pnsaDoc.add_paragraph(purposeOfPNSA3)
		
		
		disclaimer = pnsaDoc.add_heading("Disclaimer:", 1)
		disclaimerForPNSA = """	The PNSA is meant to be a proof of concept prototype. It is not meant to be used in any real capacity, but to simply show a product like this could exist in the market. It is highly advised that if there are significant concerns with your network please reach out to a certified security professional. Additionally, due to the nature of the system being built on top of Kali, please do not use this application maliciously, or to do harm onto others. The creator is not liable for any damages incurred using this prototype."""
		
		disclaimerPNSA = pnsaDoc.add_paragraph(disclaimerForPNSA)
		
		pnsaDoc.add_page_break()
		
		# check to see if each of the tests are active or not
		
		if self.cbWifi.checkState() == 2:
			aboutPNSATitle = pnsaDoc.add_heading("WiFi Analysis:", 1)
			introWifi = """	In many home and small business networks, sometimes the only security hardware present on the network is simply the router and the Wi-Fi connection. It is demandingly important to take basic precautions with your Wi-Fi Connection. Sometimes the only thing standing between you and your next breach or invasion of privacy is which encryption level you choose."""
			logDump = ""
			
			pnsaDoc.add_paragraph(introWifi)
			with open("testOutput/testAnalysis.txt", 'r') as wifiFile:
				
				logDump = wifiFile.readlines()
				# print(logDump)
				
				for line in logDump:
					#print(line)
					if "It_Hurts_When_IP" in line:
						#print(line)
						
						if "WPA2" in line:
							aboutPNSATitle = pnsaDoc.add_heading("Wi-Fi Match Found", 2)
							pnsaDoc.add_paragraph(line)
							pnsaDoc.add_paragraph("We found that your Wi-Fi uses the current highest level of encryption. This is very good, keep on doing the great job of securing your network.")
							#print("Good Level Security")
						
						elif "WPA1" in line:
							aboutPNSATitle = pnsaDoc.add_heading("Wi-Fi Match Found", 2)
							pnsaDoc.add_paragraph(line)
							pnsaDoc.add_paragraph("We found that your Wi-Fi connection is not using the highest level of encryption. The PNSA recommends upgrading to WPA2, this will secure your connection and protect your data." )
							
							#print("You should upgrade to WPA2")
						
						elif "WEP" in line:
							aboutPNSATitle = pnsaDoc.add_heading("Wi-Fi Match Found", 2)
							pnsaDoc.add_paragraph(line)
							pnsaDoc.add_paragraph("We found that your Wi-Fi connection uses WEP. The PNSA does not recommend using WEP. WEP has been found to be dangerously insecure. Please upgrade your encryption level to either WPA1 or to WPA2 if possible. This can normally be accomplished by access your router’s setting page. Please refer to your routers documentation on accessing the settings and configuration screen." )
						else:
							aboutPNSATitle = pnsaDoc.add_heading("No Wi-Fi Match Found", 2)
							pnsaDoc.add_paragraph("I was unable to find a matching WiFi Connection")
							#print("I could not detect a encryption level")	
						
						#pnsaDoc.add_paragraph(line)
			
			aboutPNSATitle = pnsaDoc.add_heading("Additional Information", 2)
			pnsaDoc.add_paragraph("	You might of noticed, that the PNSA found more then one Wi-Fi connection. There could be any number of reasons for this. Newer routers can offer a technology called dual band, this enables a faster connection to the Internet by offering two signals (one at 2.4 Ghz, and one at 5ghz). Additionally, if you have not changed the default name for your router there could be others with the same model. Finally, if the previous two statements are not true in your case, there could be someone else hosting a Wi-Fi connection to trick other into connecting to it instead of your connection. This is known as a rogue access point, this could be potentially dangerous as it could trick family members and or customers into connect to that connection and act as a man in the middle attack.")		
			
			aboutPNSATitle = pnsaDoc.add_heading("Raw Log File", 2)
			pnsaDoc.add_paragraph(logDump)
			
			pnsaDoc.add_page_break()
			
		if self.passAnal.checkState() == 2:
			aboutPNSATitle = pnsaDoc.add_heading("Password Analysis:", 1)
			introPassword = """	Insert Intro for Password Analysis"""
			logDump = ""
			
			with open("testOutput/passwordAnalysis.txt", 'r') as passwordFile:
				
				logDump = passwordFile.readlines()
				#print(logDump)
				passwordAnalysisText = "	"
				
				if "passwordShort\n" in logDump:
					passwordAnalysisText += " We found your password to be to short. A good length of a password should be at least 12-16 characters. Longer is always better."
				
				if "passwordLong\n" in logDump:
					passwordAnalysisText += " We found your password to contain enough characters. Good Job!"
		
			
				if "noDigits\n" in logDump:
					passwordAnalysisText += " We found your password does not contain any numbers. A good password should contain a few numbers that are not easy to guess."
	
				if "hasDigits\n" in logDump:
					passwordAnalysisText += " We found your password does contain at least one or more numbers. Remember the more random the numbers the harder it is to guess your password."
				
				if "noCaps\n" in logDump:
					passwordAnalysisText += " We found your password has no capital letters. A good password should include at least one capital letter to harden your password against password cracking."

				if "hasCaps\n" in logDump:
					passwordAnalysisText += " We found your password has capital letters. Your password is significantly harder to crack by using a mixture of capital and lowercase letters."
	
				if "hasSpecials\n" in logDump:
					passwordAnalysisText += " We found your password has special characters. Good Job! By using special characters this will make your password significantly harder to crack."
					
				if "noSpecials\n" in logDump:
					passwordAnalysisText += " We found your password has no special characters. It would be ideal to include at least one special character, this will make your password much harder to guess."

				if "passwordFound\n" in logDump:
					passwordAnalysisText += " We found your password on a leaked password list. We recommend changing your password as soon as possible. It is not safe to use this password with any software or hardware device."
	
				if "passwordNotFound\n" in logDump:
					passwordAnalysisText += " We did not find your password on a popular leaked password list. This does not mean that your password had not been cracked or leaked in the past, or that your password doesn't exist on other password lists. Always use caution."
			
			pnsaDoc.add_paragraph(introPassword)
			
			pnsaDoc.add_heading("Password in-depth review", 2)
			pnsaDoc.add_paragraph(passwordAnalysisText)
			
			pnsaDoc.add_heading("Raw Log File", 2)
			pnsaDoc.add_paragraph(logDump)
			
			pnsaDoc.add_page_break()
		
		if self.assDiscover.checkState() == 2:
			aboutPNSATitle = pnsaDoc.add_heading("Asset Discovery:", 1)
			with open("testOutput/liveHosts.txt", 'r') as assFile:
				logDump = assFile.readlines()
				#print(logDump)
				assAnalysisText = "	"
				assCount = 0
				for line in logDump:
					assCount+=1
				assAnalysisText += "We found that " + str(assCount) + " assets responded to a basic ping test. Ping can be used to sweep through a network and see which machines are alive. By doing this first an attacker can save time in a port scan, and reduce the amount of noise they generate so that to not waste time looking for machines that are not alive. We recommend disabling any machines from using Ping to heighten your security of the network."	
			
			introAssDiscover = """	The PNSA performed a scanning function called a port scan via a tool named NMAP. Using NMAP the PNSA can discover dangerous ports that could be leveraged to gain access to your network, steal information, and other malicious activities. It’s important to peridically scan your network to make sure attacks haven’t opened up ports you are unaware of. Below are the results of your scan."""
			
			pnsaDoc.add_paragraph(introAssDiscover)
			
			pnsaDoc.add_heading("Asset Discovery", 2)
			pnsaDoc.add_paragraph(assAnalysisText)
			
			pnsaDoc.add_heading("Raw Log File", 2)
			pnsaDoc.add_paragraph(logDump)
			
			
			pnsaDoc.add_page_break()
		
		if self.portScan.checkState() == 2:
			aboutPNSATitle = pnsaDoc.add_heading("Port Scan:", 1)
			introPortScan = """	Insert Intro for introPortScan"""
			pnsaDoc.add_paragraph(introPortScan)
			pnsaDoc.add_heading("Network Overview", 2)
			with open("testOutput/portScan.txt", 'r') as portFile:
				logDump = portFile.readlines()
				
				#print(logDump)
				
				# Start of Red Flag scanning
				
				if "22/tcp   open  ssh\n" in logDump:
					pnsaDoc.add_paragraph("Found ssh port 22 in use")
					pnsaDoc.add_paragraph("	Port 22 is generally reserved for SSH. SSH can be used to connect to a hardware device and keep the traffic secure with encryption. This isn’t a call for alarm right away, however, the PNSA recommends that if you do not use SSH to turn this service off.")
				
				if "23/tcp   open  telnet\n" in logDump:
					pnsaDoc.add_paragraph("Found telnet port 23 in use")
					pnsaDoc.add_paragraph("	Port 23 is generally reserved for Telnet. Telnet can be used to connect to a hardware device and but does not keep the data traded back and forth secured. This should be turned off immediately, most industry professionals recommend that this service is turned off. ")
				
				if "53/tcp   open  domain\n" in logDump:
					pnsaDoc.add_paragraph("Found DNS port 53 in use")
					pnsaDoc.add_paragraph("	Port 53 is generally reserved for DNS Services. DNS serves as a phonebook for your network. DNS can be used to link a hostname to an IP address. Generally for home users I might be a good idea to turn this service off if you are not using DNS for host name resolution. For a small business it’s important to harden you DNS services against zone transfers.")
					
				if "80/tcp   open  http\n" in logDump:
					pnsaDoc.add_paragraph("Found unencrypted port 80 in use")
					pnsaDoc.add_paragraph("	Port 80 is generally reserved for HTTP. If this port is open on your router, this is probably the port you would use to connect to make changes to your routers settings. Typically, you’d see this on web-servers as well. If this is on a regular user computer, PNSA recommends killing this port.")
				
				if "139/tcp  open  netbios-ssn\n" in logDump:
					pnsaDoc.add_paragraph("Found Netbios port 139 in use")
					pnsaDoc.add_paragraph("	Port 139 is generally reserved for Netbios. Most security professionals recommend for a small home network that it’s generally better to shut off Netbios services if they are not required. Many types of ransomware have been known to leverage Netbios services to spread throughout the network.")
					
				if "5000/tcp open  upnp\n" in logDump:
					pnsaDoc.add_paragraph("Found upnp (Plug and Play) port 5000 in use")
					pnsaDoc.add_paragraph("	Port 5000 is generally reserved for Plug and Play services. Generally, plug and play services are dangerous at best. The FBI even at one point recommended just shutting off all plug and play services. Due to coding errors, and general misuse of the protocol. Just don’t risk it, kill that port.")
					
				if "515/tcp  open  printer" in logDump:
					pnsaDoc.add_paragraph("Found printer port 515 in use")
					pnsaDoc.add_paragraph("	Port 515 is generally reserved for printing services. This isn’t a call for alarm immediately, however, printers can be abused in certain ways to get access to the rest of the network. Always keep an eye on which services your printer is running, and please keep the firmware up to date.")
			
			pnsaDoc.add_heading("Raw Log File", 2)
			pnsaDoc.add_paragraph(logDump)
			
			
			pnsaDoc.add_page_break()
		
		if self.dhcpAttack.checkState() == 2:
			aboutPNSATitle = pnsaDoc.add_heading("DHCP Attack Scan:", 1)
			introdhcpAttack = """	Insert Intro for DHCP Attack scan"""
			dhcpAttackInfo = ""
			
			with open("testOutput/yersinia.log", 'r') as dhcpFile:
				logDump = dhcpFile.readlines()
				#print(logDump)
				
				if "Entering command line mode...\n" in logDump:
					dhcpAttackInfo = """	We 'Tested' the water for what we believe is your router to see if it could succumb to a DHCP attack. Based on the limited tests, we were able to determine that there is a possibility your router could be knocked offline to Denial Of Service attacks. We recommend changing your router settings to only lease up to 50 DHCP leases, or only approve certain MAC addresses."""
				
			
			
			
			
			pnsaDoc.add_paragraph(introdhcpAttack)
			
			pnsaDoc.add_heading("DHCP Attack Review", 2)
			pnsaDoc.add_paragraph(dhcpAttackInfo)
			
			pnsaDoc.add_heading("Raw Log File", 2)
			pnsaDoc.add_paragraph(logDump)
			
			pnsaDoc.add_page_break()
		
		
		# too save to a location do this 
		# filename = "path/to/save/location/" + nameoffile
		# pnsaDoc.save(filename)
		
		# too conver to any format use...
		# soffice --headless --convert-to pdf reportPDFTEST.docx
		# pdf, rtf, txt work this function will default to docx, then do conversion if needed
		
		pnsaDoc.save('PNSA_Report.docx')
		
		PNSA_Report_File = "PNSA_Report."
		#As the document gets wordier more time needs to be given to finish the write out operations
		time.sleep(5)
		
		if selectedType == "TXT":
			command = "soffice --headless --convert-to txt PNSA_Report.docx"
			os.system(command)
			time.sleep(10)
			TXTSRC = PNSA_Report_File + "txt"
			TXTDST = self.saveLocationDir + TXTSRC
			copyfile(TXTSRC, TXTDST)
			# Start cleanup
			
		elif selectedType == "PDF":
			command = "soffice --headless --convert-to pdf PNSA_Report.docx"
			os.system(command)
			time.sleep(10)
			PDFSRC = PNSA_Report_File + "pdf"
			PDFDST = self.saveLocationDir + "/" + PDFSRC
			copyfile(PDFSRC, PDFDST)
			time.sleep(10)
			# Start cleanup
			
		elif selectedType == "RTF":
			command = "soffice --headless --convert-to rtf PNSA_Report.docx"
			os.system(command)
			time.sleep(10)
			RTFSRC = PNSA_Report_File + "rtf"
			RTFDST = self.saveLocationDir + "/" + RTFSRC
			copyfile(RTFSRC, RTFDST)
			time.sleep(10)
			# Start cleanup
			
		else:
			time.sleep(10)
			SRC = PNSA_Report_File + ".docx"
			DST = self.saveLocationDir + "/" + SRC
			copyfile(SRC, DST)
			time.sleep(10)
			
			# Start cleanup
	def testingSequence(self):
		print(self.nameInputText)	
		print(self.wifiNameText)
		print(self.wifiPasswordText)
		print(self.portScan.checkState())
		
		# Get IP address
		# Kali uses eth0 (ether) wlan0 (wireless) below is for my laptop tests
		self.ipAddressUsed = networkInterface.ifaddresses('wlan0')[networkInterface.AF_INET][0]['addr']
		print(self.ipAddressUsed)
		
		
		
		if self.cbWifi.checkState() == 2:
			#This will need to be changed to fit the needs of Kali (Linux Mint for my laptop)
			command = "nmcli dev wifi list > /root/Desktop/UI_GoodProd/testOutput/testAnalysis.txt"
			print("Starting Wifi Analysis")
			os.system(command)
			
			
		if self.passAnal.checkState() == 2:
			yourPassword = self.wifiPasswordText
			isPasswordFound = False
			print("Starting Password Analysis")
			print("Your Password is: " + yourPassword)
			print("Your Password is: " + str(len(yourPassword)) + " characters long")
			
			# if false will return None
			#print(re.search(r'\d', yourPassword))
			#print(re.search('[A-Z]', yourPassword))
			#print(re.search('[!@#$%^&*()_+]', yourPassword))
			
			containsDigits = re.search(r'\d', yourPassword)
			containsCaps = re.search('[A-Z]', yourPassword)
			containsSpecials = re.search('[!@#$%^&*()_+]', yourPassword)
			
			
			if len(yourPassword) < 12:
				print("your password is too short")
				
			# Checks password against rockyou list to see if it exists	
			# had to use latin-1 encoding to get all the latin character if you use 
			# UTF-8 it freaks out and abors
			with open('PNSA_LISTS/rockyou.txt', 'r', encoding = 'latin-1') as rockYouList:
				#for line in rockYouList:
				if yourPassword in rockYouList.read():
					print("Found it")
					isPasswordFound = True 
				else:
					print("I did not find your password")
			
			# Need to do a text dump for report parsing later
			command = "touch /root/Desktop/UI_GoodProd/testOutput/passwordAnalysis.txt"
			os.system(command)
			
			# Calculating how fast it can be cracked maybe. 
			
			#Worst Case Scenerio
			if (containsDigits == None) and (containsCaps == None) and (containsSpecials == None):
				totalCombination = 26**len(yourPassword)
				# This is number of tries a second
				# Based on a 1080 looks like it can do 1.4 Billion hashes a sec with SHA1
				
				
				averageCrackGPU = 14000000000
				
				crackedIn = totalCombination / averageCrackGPU
				
				day = crackedIn // (24*3600)
				remainingSecs = crackedIn % (24*3600)
				
				
				hours = remainingSecs // 3600
				remainingSecs %= 3600
				
				mins = remainingSecs // 60
				remainingSecs %= 60
				
				secs = remainingSecs
				
				print("This could be cracked in... ")
				print("Days: %s \n Hours: %d \n Minutes: %d \n Seconds: %f" % ("{:,}".format(day), hours, mins, secs))

			#Best Case Scenerio	
			else:
				totalCombination = 72**len(yourPassword)
				averageCrackGPU = 14000000000
				
				crackedIn = totalCombination / averageCrackGPU
				
				day = crackedIn // (24*3600)
				remainingSecs = crackedIn % (24*3600)
				
				
				hours = remainingSecs // 3600
				remainingSecs %= 3600
				
				mins = remainingSecs // 60
				remainingSecs %= 60
				
				secs = remainingSecs
				
				print("This could be cracked in... ")
				print("Days: %s \n Hours: %d \n Minutes: %d \n Seconds: %f" % ("{:,}".format(day), hours, mins, secs))
				
			# All this doesn't take into consideration rainbow tables, password breaches, social engineering, ect	
			
			
			with open('testOutput/passwordAnalysis.txt', 'a') as fileOut:
				
				fileOut.write("PASSWORD ANALYSIS\n\n")
				fileOut.write("Password: " + yourPassword + "\n")
				fileOut.write("Password Length: " + str(len(yourPassword)) + "\n")
				
				if len(yourPassword) < 12:
					fileOut.write("passwordShort\n")
				else:
					fileOut.write("passwordLong\n")
				
				
				if containsDigits != None:
					fileOut.write("hasDigits\n")
				else:
					fileOut.write("noDigits\n")
				
				if containsCaps != None:
					fileOut.write("hasCaps\n")
				else:
					fileOut.write("noCaps\n")
				
				if containsSpecials != None:
					fileOut.write("hasSpecials\n")
				else:
					fileOut.write("noSpecials\n")
				
				if isPasswordFound == True:
					fileOut.write("passwordFound\n")
				else:
					fileOut.write("passwordNotFound\n")
					
				#for i in range(11):
				#	fileOut.write("Test line %d\r\n" % (i+1))
			
		if self.assDiscover.checkState() == 2:
			
			
			currentIP = self.ipAddressUsed
			cutIPAddr = currentIP.split(".", 3)[:3]
			
			
			IPHostsubNet = ("" + cutIPAddr[0] + "." + cutIPAddr[1] + "." + cutIPAddr[2] + '.1')
			self.probablyRouter = IPHostsubNet
			
			IPHostsubNet = IPHostsubNet + "/24"
			print("Your subnet is probably " + IPHostsubNet)
			print("Starting Asset Discovery")
			command = "fping -g -a " + IPHostsubNet  + " > /root/Desktop/UI_GoodProd/testOutput/liveHosts.txt"
			os.system(command)
			
			
		if self.portScan.checkState() == 2:
			print("Starting Port Scanning")
			
			currentIP = self.ipAddressUsed
			cutIPAddr = currentIP.split(".", 3)[:3]
			
			
			IPHostsubNet = ("" + cutIPAddr[0] + "." + cutIPAddr[1] + "." + cutIPAddr[2] + '.1')
			
			IPHostsubNet += "/24"
			
			print(IPHostsubNet)
			command ="sudo nmap -O -F " + IPHostsubNet + " > /root/Desktop/UI_GoodProd/testOutput/portScan.txt"
			
			os.system(command)
		
		
		if self.dhcpAttack.checkState() == 2:
			print("Starting DHCP Starvation")
			
			
			# yersinia dhcp -attack 1 -source b8:27:eb:of:a9:10 -dest E4:F4:C6:10:4E:39 -interface wlan0
			
			currentIP = self.ipAddressUsed
			cutIPAddr = currentIP.split(".", 3)[:3]
			
			
			IPHostCurr = ("" + cutIPAddr[0] + "." + cutIPAddr[1] + "." + cutIPAddr[2] + '.1')
			
			self.probablyRouter = IPHostCurr
			
			commandARP = "arping -I wlan0 -c 1 " + self.probablyRouter + " > /root/Desktop/UI_GoodProd/testOutput/arp.txt"
			commandARP2 = "arping -I wlan0 -c 1 " + self.probablyRouter
			
			
			os.system(commandARP)
			routerMac = ""
			
			with open('testOutput/arp.txt', 'r') as readFile:
				arpToSearch = readFile.read()
				
				mac = re.search(r"((?:[0-9a-fA-F]:?){12})", arpToSearch)
				routerMac = mac.group()
			
			#commandDHCP = "dhcpstarv -v -d " + routerMac + " -i wlp5s0 > /root/Desktop/UI_GoodProd/testOutput/dhcp.txt"
			
			commandDHCP = "yersinia dhcp -attack 1 -source b8:27:eb:of:a9:10 -dest " + routeMac + " -interface wlan0"
			
			P = subprocess.Popen(commandDHCP,shell=True)
			time.sleep(10)
			P.terminate()

			# Move the log into the correct folder section
			
			copyfile("/root/yersinia.log", "/root/Desktop/UI_GoodProd/testOutput/yersinia.log")
			
# Standard Main Loop That Everyone uses for PyQt5 
if __name__ == '__main__':
	app = QApplication(sys.argv)
	app.setStyle(QStyleFactory.create("GTK+"))
	ex = App()
	ex.show()
	sys.exit(app.exec_())
